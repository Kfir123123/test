import os
import docx
import datetime

from docx.api import Document
from fpdf import FPDF

VERIFICATION_TEST_PATH = r"C:\1.22"
OUTPUT_PATH = ""
TESTER_SIGNATURE = r"C:\Users\admin\Desktop\signature.JPEG"
MAX_LINE_LENGTH = 80


def write_text_table(pdf, cell_width, cell_height, lines, move_to_next_line):
    for line in lines:
        pdf.cell(cell_width, cell_height, txt=line.replace('_', ''), ln=move_to_next_line)


def split_text_table(text, additional_space):
    lines = []
    add_space = ""

    if len(text) <= MAX_LINE_LENGTH:
        lines.append(text)
    else:
        i = 0
        while i in range(0, len(text)):
            if i + MAX_LINE_LENGTH < len(text):
                current_line = add_space + text[i:i + MAX_LINE_LENGTH]
            else:
                current_line = add_space + text[i:len(text)]

            add_space = additional_space

            i += MAX_LINE_LENGTH
            lines.append(current_line)

    return lines


def write_text(pdf, lines):
    for line in lines:
        pdf.cell(0, 10, txt=line.replace('_', ''), ln=1)


def split_text_helper(text, lines):
    i = 0
    while i in range(0, len(text)):
        if i + MAX_LINE_LENGTH < len(text):
            current_line = text[i:i + MAX_LINE_LENGTH]
        else:
            current_line = text[i:len(text)]

        i += MAX_LINE_LENGTH
        lines.append(current_line)


def split_text(text):
    lines = []
    flag = False
    if "Configurations" and "Files" in text:
        split_text = text.split("Files")
        split_text_helper(split_text[0], lines)
        split_text_helper("Files" + split_text[1], lines)
        flag = True

    elif len(text) <= MAX_LINE_LENGTH:
        lines.append(text)
    else:
        split_text_helper(text, lines)

    if flag:
        new_lines = []
        for line in lines:
            if len(line) > MAX_LINE_LENGTH:
                new_lines.append(line[:MAX_LINE_LENGTH])
                new_lines.append(line[MAX_LINE_LENGTH:])
            else:
                new_lines.append(line)

    return lines


def pre_conditions(file_path):
    doc = docx.Document(file_path)
    paragraphs = []
    for paragraph in doc.paragraphs:
        paragraphs.append(paragraph.text)

    for line in paragraphs:
        if "Pre conditions" in line:
            return line
        if "Preconditions" in line:
            return line.replace("Preconditions", "Pre conditions")

    return "Couldn't find the Pre conditions instructions"


def read_table():
    keys = None
    for i, row in enumerate(table.rows):
        text = []
        for cell in row.cells:
            text.append(cell.text)

        # Establish the mapping based on the first row
        # headers; these will become the keys of our dictionary
        if i == 0:
            keys = tuple(text)
            continue

        # Construct a dictionary for this row, mapping
        # keys to values for this row
        row_data = dict(zip(keys, text))
        data.append(row_data)


def generate_pdf(docx_path, pdf_path):
    # Open the DOCX file
    doc = Document(docx_path)

    # Create a PDF object
    pdf = FPDF('P', 'mm', 'A4')
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()

    first_time = True
    write_table = False

    # Iterate over the paragraphs in the document
    for paragraph in doc.paragraphs:
        paragraph.text.replace('_', '')
        if first_time:
            pdf.set_font("Arial", size=15, style="B")
            for line in split_text(paragraph.text):
                pdf.cell(0, 10, txt=line, ln=1, align="C")
            first_time = False
            continue

        pdf.set_font("Arial", size=12)

        if "Tester" in paragraph.text:
            x = pdf.get_x() + 150
            y = pdf.get_y() - 5
            pdf.cell(0, 10, txt="Tester: " + tester_name + "\t" * 75 + "Signature: ", ln=1)

            pdf.image(TESTER_SIGNATURE, x=x, y=y, w=30, h=20)  # Add the signature image
            continue

        elif "Date" in paragraph.text:
            pdf.cell(0, 10, txt="Date: " + datetime.datetime.now().strftime("%Y-%m-%d"), ln=1)
            continue

        elif "version" in paragraph.text:
            pdf.cell(0, 10, txt="SW version \ build number: " + version_number, ln=1)
            continue

        elif "Summary" in paragraph.text:
            write_text(pdf, split_text("Summary, conclusion and recommendations: " + summary))
            continue

        elif write_table:
            pdf.set_font("Arial", size=10)
            first_row = True
            test_info = True
            result_info = False
            result = False
            index = 1

            for row_data in data:
                for head_line, to_print in row_data.items():
                    # skip first row
                    if first_row:
                        first_row = False
                        continue

                    # first col
                    if test_info:
                        if to_print == "":
                            first_row = True
                            test_info = True
                            continue
                        write_text_table(pdf, 0, 7, split_text_table(str(index) + ". Test: " + to_print, "    "), 1)
                        test_info = False
                        result_info = True

                    # second col
                    elif result_info:
                        write_text_table(pdf, 0, 7, split_text_table("    Expected_Results: " + to_print, "       "), 1)
                        result_info = False
                        result = True


                    elif result:
                        pdf.cell(40, 7, txt="    Result:", ln=0)

                        if index <= len(test_results):
                            # Set the text color based on the result
                            if test_results[index - 1] == "v":
                                pdf.set_text_color(0, 128, 0)  # Set text color to green for "Pass"
                                pdf.cell(35, 7, txt="Pass", ln=1, align='L')
                            else:
                                pdf.set_text_color(255, 0, 0)  # Set text color to red for "Fail"
                                pdf.cell(35, 7, txt="Fail", ln=1, align='L')

                        # Reset the text color to black
                        pdf.set_text_color(0, 0, 0)
                        result = False

                    # fourth col
                    else:
                        if index <= len(comment_results) and comment_results[index - 1] != "":
                            write_text_table(pdf, 0, 7,
                                             split_text_table("    Comments: " + comment_results[index - 1], "       "),
                                             1)

                        first_row = True
                        test_info = True
                        index += 1

            pdf.ln()
            write_table = False

        if "Pre conditions" in paragraph.text:
            write_table = True

        write_text(pdf, split_text(paragraph.text))

    # Save the PDF file
    with open(pdf_path, 'wb') as file:
        file.write(pdf.output(dest='S').encode('latin-1'))


for test_type in os.listdir(VERIFICATION_TEST_PATH):
    print(test_type.upper())
    version_number = input("please insert the version: ")

    for test_file in os.listdir(os.path.join(VERIFICATION_TEST_PATH, test_type)):
        if not test_file.endswith("docx"):
            continue

        cont = ""
        for test_file_temp in os.listdir(os.path.join(VERIFICATION_TEST_PATH, test_type)):
            if test_file.replace('.docx', '') in test_file_temp and "pdf" in test_file_temp:
                if test_file.replace('.docx', '') == test_file_temp.replace(".docx", ''):
                    cont = input(f'this test case {test_file} has already been tested, would you like to skip? y/n ')
                    while cont != 'y' and cont != 'n':
                        cont = input(
                            'Illegal input \nthis test case has already been tested, would you like to skip? y/n ')
                        if cont == 'y':
                            break
                    break

        if cont == 'y':
            continue

        while True:
            tester_name = input("please insert tester name: ")
            if tester_name != "":
                break

        test_path = os.path.join(VERIFICATION_TEST_PATH, test_type, test_file)

        date_string = datetime.datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
        pdf_name = test_file.replace(".docx", f"_{date_string}.pdf")
        OUTPUT_PATH = os.path.join(VERIFICATION_TEST_PATH, test_type, pdf_name)

        print(f"\nFile name: {test_file}")
        print(pre_conditions(test_path) + "\n")

        # save input
        test_results = []
        comment_results = []
        idx = 0

        # read table
        document = Document(test_path)
        table = document.tables[0]

        data = []
        read_table()

        first_row = True
        test_info = True
        result_info = False
        result = False
        break_flag = False

        for idx, row_data in enumerate(data):
            if break_flag:
                break
            for head_line, to_print in row_data.items():
                # skip first row
                if first_row:
                    first_row = False
                    continue

                if "Test" in head_line and to_print == "":
                    break_flag = True
                    break
                print(f"{head_line}:", end=" ")

                # first col
                if test_info:
                    if to_print == "":
                        first_row = True
                        test_info = True
                        continue

                    print(f"{to_print}")
                    test_info = False
                    result_info = True

                # second col
                elif result_info:
                    print(f"{to_print}")
                    result_info = False
                    result = True

                # third col
                elif result:
                    while True:
                        res = input("v/x? ").lower()
                        if res in ['v', 'x']:
                            test_results.append(res)
                            result = False
                            break
                        else:
                            print("Choose 'v' or 'x'")

                # fourth col
                else:
                    print("(press enter to continue)")
                    comment = input().lower()
                    comment_results.append(comment)
                    first_row = True
                    test_info = True
                    idx += 1

        summary = input("Summary, conclusion and recommendations: (press enter to continue)")
        generate_pdf(test_path, OUTPUT_PATH)
